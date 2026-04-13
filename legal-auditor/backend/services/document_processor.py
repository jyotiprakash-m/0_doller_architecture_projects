"""
Document processor — extracts text from PDF, DOCX, and TXT files.
All processing happens locally. No external API calls.
"""
import os
from pathlib import Path
from typing import Optional


def extract_text_from_pdf(file_path: str) -> tuple[str, int]:
    """Extract text from a PDF file. Returns (text, page_count)."""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages), len(reader.pages)


def extract_text_from_docx(file_path: str) -> tuple[str, int]:
    """Extract text from a DOCX file. Returns (text, paragraph_count)."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs), len(paragraphs)


def extract_text_from_txt(file_path: str) -> tuple[str, int]:
    """Extract text from a TXT file. Returns (text, line_count)."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    line_count = len(text.splitlines())
    return text, line_count


def extract_text(file_path: str, file_type: str) -> tuple[str, int]:
    """
    Extract text from a file based on its type.
    Returns (extracted_text, page_or_section_count).
    """
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
    }

    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    text, count = extractor(file_path)

    if not text.strip():
        raise ValueError(f"No text content could be extracted from {file_path}")

    return text, count


def chunk_text(text: str, chunk_size: int = 1024, chunk_overlap: int = 200) -> list[str]:
    """
    Split text into overlapping chunks for embedding.
    Uses sentence-aware splitting to avoid cutting mid-sentence.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    sentences = text.replace("\n\n", "\n").split(". ")
    current_chunk = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        if not sentence.endswith("."):
            sentence += "."

        if len(current_chunk) + len(sentence) + 1 <= chunk_size:
            current_chunk += (" " + sentence) if current_chunk else sentence
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
                # Overlap: keep the tail of the current chunk
                overlap_text = current_chunk[-chunk_overlap:] if len(current_chunk) > chunk_overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
            else:
                # Single sentence exceeds chunk size — split by characters
                chunks.append(sentence[:chunk_size])
                current_chunk = sentence[chunk_size - chunk_overlap:]

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


def save_uploaded_file(file_content: bytes, filename: str, upload_dir: str) -> str:
    """Save uploaded file to the data directory."""
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, filename)

    # Handle duplicate filenames
    if os.path.exists(file_path):
        name, ext = os.path.splitext(filename)
        counter = 1
        while os.path.exists(file_path):
            file_path = os.path.join(upload_dir, f"{name}_{counter}{ext}")
            counter += 1

    with open(file_path, "wb") as f:
        f.write(file_content)

    return file_path
